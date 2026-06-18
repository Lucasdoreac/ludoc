import json, urllib.request, urllib.error, http.server, subprocess, os, sqlite3, uuid, threading, datetime, html.parser, urllib.parse, re, ast, docx, shlex, socket
import llm, cleaner, acp_client

# ... (rest of imports) ...

def _generate_brand_docs(params):
    template_name = params.get("template_name")
    data = params.get("data", {})
    if not template_name or not isinstance(data, dict):
        return 400, {"error": "template_name and data required"}
    
    template_path = os.path.join(WORKSPACE, "brand_assets", template_name)
    output_dir = os.path.join(WORKSPACE, "deliverables")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"{uuid.uuid4()}.docx")
    
    try:
        doc = docx.Document(template_path)
        for p in doc.paragraphs:
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))
        doc.save(output_path)
        return 200, {"status": "ok", "path": output_path}
    except Exception as e:
        return 500, {"error": str(e)}

WORKSPACE  = "/tmp/workspace"
DB_PATH    = os.path.join(WORKSPACE, "tasks.db")
PROXY      = os.environ.get("SKILL_PROXY", "http://localhost:8000")

# ACP Configuration
ACP_REGISTRY = {}  # agent_id -> {url, capabilities, last_seen}
ACP_AGENT_ID = os.environ.get("ACP_AGENT_ID", f"ludoc-{uuid.uuid4().hex[:8]}")
ACP_CAPABILITIES = []  # Populated from skills catalog

os.makedirs(WORKSPACE, exist_ok=True)

# --- Banco de Dados ---
def init_db():
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS jobs (
                task_id TEXT PRIMARY KEY,
                status TEXT,
                prompt TEXT,
                result TEXT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)

# --- Carga de Skills ---
SKILLS_DIR = os.path.join(os.path.dirname(__file__), "skills")
_catalog = {}

def _load_skills():
    global _catalog
    _catalog = {}
    if not os.path.exists(SKILLS_DIR): return
    for root, dirs, files in os.walk(SKILLS_DIR):
        if "skill.md" in files:
            path = os.path.join(root, "skill.md")
            try:
                with open(path, "r", encoding="utf-8") as f: content = f.read()
                match = re.search(r"^---\s*\n(.*?)\n---\s*\n", content, re.DOTALL)
                if match:
                    meta = {}
                    for line in match.group(1).split("\n"):
                        if ":" in line:
                            k, v = line.split(":", 1)
                            k, v = k.strip(), v.strip()
                            if v.startswith("[") and v.endswith("]"): v = [x.strip().strip("'").strip('"') for x in v[1:-1].split(",") if x.strip()]
                            meta[k] = v
                    if "name" in meta:
                        print(f"DEBUG: Loaded skill {meta['name']} with upstream: {meta.get('upstream')}")
                        meta["system_instruction"] = content[match.end():].strip()
                        _catalog[meta["name"]] = meta
            except Exception as e: print(f"Erro ao carregar skill em {path}: {e}")
_load_skills()

def _init_acp():
    global ACP_CAPABILITIES
    ACP_CAPABILITIES = [
        {"name": s["name"], "description": s["description"], "params": s["params"]}
        for s in _catalog.values()
    ]
    print(f"ACP Agent {ACP_AGENT_ID} initialized with {len(ACP_CAPABILITIES)} capabilities")
_init_acp()

def list_skills():
    return [{"name": s["name"], "description": s["description"], "params": s["params"]} for s in _catalog.values()]

# --- Handlers ---
_CWD = [WORKSPACE]

def _exec_shell(params):
    cmd      = params.get("command", "")
    timeout  = int(params.get("timeout", 30))
    approved = params.get("approved", False)
    if not cmd: return 400, {"error": "command required"}
    blacklist = ["rm", "push", "kill", "wget", "curl", "apt", "sh", "bash"]
    if not approved and any(re.search(rf"\b{term}\b", cmd) for term in blacklist):
        return 200, {"error": "COMANDO BLOQUEADO: Risco de segurança. Requer HITL (Aprovação Humana) para execução"}
    stripped = cmd.strip()
    args = shlex.split(stripped)
    if not args: return 400, {"error": "command required"}
    
    if args[0] == "cd":
        if len(args) < 2: return 200, {"stdout": "", "stderr": "", "code": 0, "cwd": _CWD[0]}
        target = args[1]
        new_cwd = target if os.path.isabs(target) else os.path.normpath(os.path.join(_CWD[0], target))
        if os.path.isdir(new_cwd):
            _CWD[0] = new_cwd
            return 200, {"stdout": "", "stderr": "", "code": 0, "cwd": _CWD[0]}
        return 200, {"stdout": "", "stderr": f"cd: {target}: No such directory", "code": 1}
    try:
        res = subprocess.run(args, capture_output=True, text=True,
                             cwd=_CWD[0], timeout=60)

        LIMIT = 30000
        stdout = res.stdout
        stderr = res.stderr
        truncated = False
        if len(stdout) > LIMIT:
            stdout = stdout[:15000] + f"\n...[truncated {len(stdout)-LIMIT} chars]...\n" + stdout[-5000:]
            truncated = True
        return 200, {"stdout": stdout, "stderr": stderr[:5000], "code": res.returncode, "cwd": _CWD[0], "truncated": truncated}
    except subprocess.TimeoutExpired:
        return 200, {"stdout": "", "stderr": f"timeout after {timeout}s", "code": -1, "cwd": _CWD[0]}

def _fs_read(params):
    path   = params.get("file_path", "")
    offset = int(params.get("offset", 0))
    limit  = int(params.get("limit", 500))
    if not path: return 400, {"error": "file_path required"}
    full = path if os.path.isabs(path) else os.path.join(_CWD[0], path)
    safe = os.path.normpath(full)
    try:
        with open(safe, "r", errors="replace") as f:
            lines = f.readlines()
        total = len(lines)
        chunk = lines[offset : offset + limit]
        content = "".join(f"{offset+i+1:4d} | {l}" for i, l in enumerate(chunk))
        return 200, {"path": safe, "content": content, "total_lines": total, "offset": offset, "limit": limit, "partial": (offset+limit) < total}
    except Exception as e: return 500, {"error": str(e)}

def _fs_search(params):
    pattern = params.get("pattern", "")
    path    = params.get("path", WORKSPACE)
    if not pattern: return 400, {"error": "pattern required"}
    full = path if os.path.isabs(path) else os.path.join(WORKSPACE, path)
    try:
        res = subprocess.run(["grep", "-rnE", pattern, "."], capture_output=True, text=True, cwd=full, timeout=15)
        return 200, {"stdout": res.stdout[:10000], "stderr": res.stderr}
    except Exception as e: return 500, {"error": str(e)}

def _python_interpreter(params):
    code = params.get("code", "")
    if not code: return 400, {"error": "code required"}
    try:
        res = subprocess.run(["python3", "-c", code], capture_output=True, text=True, timeout=10)
        return 200, {"stdout": res.stdout, "stderr": res.stderr, "code": res.returncode}
    except Exception as e: return 500, {"error": str(e)}

def _write_file(params):
    filename = params.get("filename", "")
    content  = params.get("content", "")
    if not filename: return 400, {"error": "filename required"}
    safe = os.path.normpath(os.path.join(WORKSPACE, filename))
    if not safe.startswith(WORKSPACE): return 403, {"error": "path traversal denied"}
    parent = os.path.dirname(safe)
    if parent: os.makedirs(parent, exist_ok=True)
    with open(safe, "w") as f: f.write(content)
    return 200, {"status": "ok", "path": safe}

def _fs_edit(params):
    file_path  = params.get("file_path", "")
    old_string = params.get("old_string", "")
    new_string = params.get("new_string", "")
    if not file_path or old_string == "": return 400, {"error": "file_path and old_string required"}
    full = file_path if os.path.isabs(file_path) else os.path.join(_CWD[0], file_path)
    safe = os.path.normpath(full)
    try:
        with open(safe, "r", errors="replace") as f: content = f.read()
        count = content.count(old_string)
        if count == 0: return 422, {"error": "old_string not found"}
        if count > 1: return 422, {"error": f"matches {count} locations"}
        new_content = content.replace(old_string, new_string, 1)
        with open(safe, "w") as f: f.write(new_content)
        return 200, {"status": "ok", "path": safe, "replacements": 1}
    except Exception as e: return 500, {"error": str(e)}

_EPISODES_LOCK = threading.Lock()
def _distill_experience(params):
    task = params.get("task", "").replace("\n", " ").replace("|", "-").strip()
    solution = params.get("solution", "").replace("\n", " ").replace("|", "-").strip()
    if not task or not solution: return 400, {"error": "task/solution required"}
    ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    line = f"[{ts}] Tarefa: {task} | Solução: {solution}\n"
    path = "/tmp/workspace/episodes.md"
    try:
        with _EPISODES_LOCK:
            with open(path, "a", encoding="utf-8") as f: f.write(line)
        return 200, {"status": "ok"}
    except Exception as e: return 500, {"error": str(e)}

def _validate_code(params):
    code = params.get("code", "")
    if not code: return 400, {"error": "code required"}
    try:
        ast.parse(code)
    except SyntaxError as e: return 200, {"valid": False, "error": str(e)}
    tmp = os.path.join(WORKSPACE, "_validate_tmp.py")
    with open(tmp, "w") as f: f.write(code)
    try:
        res = subprocess.run(["ruff", "check", "--select=E,F", tmp], capture_output=True, text=True)
        os.unlink(tmp)
        if res.returncode != 0: return 200, {"valid": False, "error": res.stdout.strip()}
    except Exception:
        if os.path.exists(tmp): os.unlink(tmp)
    return 200, {"valid": True}

def _analyze_dependency_graph(params):
    nodes = list(_catalog.keys()); edges = []
    try:
        with open(__file__) as f: tree = ast.parse(f.read())
        for node in ast.walk(tree):
            if isinstance(node, ast.FunctionDef) and node.name.startswith("_"):
                skill_name = node.name.lstrip("_")
                if skill_name in _catalog:
                    for child in ast.walk(node):
                        if isinstance(child, ast.Call) and isinstance(child.func, ast.Attribute):
                            edges.append({"from": skill_name, "calls": child.func.attr})
    except Exception: pass
    return 200, {"nodes": nodes, "edges": edges}

def _web_fetch(params):
    url = params.get("url", "")
    if not url: return 400, {"error": "url required"}
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=15) as r: raw = r.read().decode(errors="replace")
        class _S(html.parser.HTMLParser):
            def __init__(self): super().__init__(); self.parts = []; self._skip = False
            def handle_starttag(self, t, a): self._skip = t in ("script","style")
            def handle_endtag(self, t): self._skip = False
            def handle_data(self, d):
                if not self._skip and d.strip(): self.parts.append(d.strip())
        p = _S(); p.feed(raw)
        return 200, {"url": url, "text": "\n".join(p.parts)[:8000]}
    except Exception as e: return 500, {"error": str(e)}

def _web_search(params):
    query = params.get("query", "")
    if not query: return 400, {"error": "query required"}
    try:
        url = "https://lite.duckduckgo.com/lite/?q=" + urllib.parse.quote_plus(query)
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=15) as r: raw = r.read().decode(errors="replace")
        results = []
        links = re.findall(r'uddg=([^&"]+)', raw); titles = re.findall(r"class='result-link'>(.*?)</a>", raw)
        for i in range(min(5, len(titles))):
            results.append({"url": urllib.parse.unquote(links[i]), "title": re.sub(r"<[^>]+>", "", titles[i]).strip()})
        return 200, {"query": query, "results": results}
    except Exception as e: return 500, {"error": str(e)}

def _git(params):
    cmd = params.get("command", "")
    if not cmd: return 400, {"error": "command required"}
    blocked = ["push", "reset --hard", "clean -f"]
    if any(b in cmd for b in blocked): return 403, {"error": "blocked"}
    try:
        git_cmd = ["git"] + cmd.split()
        res = subprocess.run(git_cmd, capture_output=True, text=True, cwd=WORKSPACE, timeout=15)
        return 200, {"stdout": res.stdout, "stderr": res.stderr, "code": res.returncode}
    except Exception as e: return 500, {"error": str(e)}

def _last30days(params):
    topic = params.get("topic", "")
    if not topic: return 400, {"error": "topic required"}
    try:
        engine_path = os.environ.get("LAST30DAYS_ENGINE", os.path.join(SKILLS_DIR, "last30days", "scripts", "last30days.py"))
        res = subprocess.run(["python3", engine_path, topic], capture_output=True, text=True, timeout=120)
        return 200, {"stdout": res.stdout, "stderr": res.stderr, "code": res.returncode}
    except Exception as e: return 500, {"error": str(e)}

def _generate_brand_docs(params):
    template_name = params.get("template_name")
    data = params.get("data", {})
    if not template_name or not isinstance(data, dict): return 400, {"error": "template_name/data required"}
    
    template_path = os.path.join(WORKSPACE, "brand_assets", template_name)
    output_dir = os.path.join(WORKSPACE, "deliverables")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"{uuid.uuid4()}.docx")
    
    try:
        import docx
        doc = docx.Document(template_path)
        for p in doc.paragraphs:
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in p.text: p.text = p.text.replace(placeholder, str(value))
        doc.save(output_path)
        return 200, {"status": "ok", "path": output_path}
    except Exception as e: return 500, {"error": str(e)}

def _bridge_mcp(params):
    command = params.get("command", "")
    if not command: return 400, {"error": "command required"}
    try:
        # Simula a ponte enviando via proxy para o host
        res = subprocess.run(shlex.split(command), capture_output=True, text=True, timeout=120)
        return 200, {"stdout": res.stdout, "stderr": res.stderr, "code": res.returncode}
    except Exception as e: return 500, {"error": str(e)}

def _security_audit(params):
    skill_path = params.get("skill_path", "")
    if not skill_path: return 400, {"error": "skill_path required"}
    try:
        res = subprocess.run(["skillspector", "scan", skill_path, "--format", "json"], 
                             capture_output=True, text=True, timeout=300)
        data = json.loads(res.stdout)
        # Lógica de governança: falha se score de risco for > 50
        risk_score = data.get("risk_score", 0)
        status = "PASSED" if risk_score <= 50 else "FAILED"
        return 200, {"status": status, "risk_score": risk_score, "details": data}
    except Exception as e: return 500, {"error": str(e)}

_INTERNAL = {
    "execute_shell": _exec_shell, "fs_read": _fs_read, "fs_search": _fs_search,
    "python_interpreter": _python_interpreter, "write_file": _write_file,
    "fs_edit": _fs_edit, "distill_experience": _distill_experience,
    "validate_code": _validate_code, "analyze_dependency_graph": _analyze_dependency_graph,
    "web_fetch": _web_fetch, "web_search": _web_search, "git": _git, "last30days": _last30days,
    "generate_brand_docs": _generate_brand_docs, "bridge_mcp": _bridge_mcp,
    "security_audit": _security_audit
}

# --- Motor Assíncrono (Agent Worker) ---
def agent_worker(task_id, prompt, caller):
    print(f"DEBUG: Starting agent_worker for {task_id}, WORKSPACE: {WORKSPACE}", flush=True)
    try:
        mem_context = ""
        for f in ["SKILL.md", "CONTEXT.md"]:
            p = os.path.join(os.path.dirname(__file__), "..", "..", f)
            if os.path.exists(p):
                with open(p, "r", encoding="utf-8") as file: mem_context += f"\n### {f}:\n{file.read()}\n"
        ep = "/tmp/workspace/episodes.md"
        if os.path.exists(ep):
            with open(ep, "r", encoding="utf-8") as file: mem_context += f"\n### Episodes:\n{file.read()}\n"
        
        relevant = cleaner.select(prompt, _catalog)
        ctx = cleaner.format_for_llm(relevant)
        msgs = [{"role": "user", "content": prompt}]
        trace = []; last = None
        
        for i in range(10): # max_iter
            dec = llm.call(llm.trim_history(msgs), model=llm.THINKER_MODEL, tools_context=ctx + "\n" + mem_context)
            print(f"DEBUG: LLM decision: {dec}", flush=True)
            trace.append({"iter": i+1, "decision": dec})
            act = dec.get("action", "final_answer"); params = dec.get("params", {})
            if (act, json.dumps(params)) == last and act != "final_answer": break
            last = (act, json.dumps(params))
            if act == "final_answer":
                _distill_experience({"task": prompt[:500], "solution": params.get('text','')[:500]})
                with sqlite3.connect(DB_PATH) as conn:
                    conn.execute("UPDATE jobs SET status = 'completed', result = ? WHERE task_id = ?", 
                                 (json.dumps({"answer": params.get('text',''), "trace": trace}), task_id))
                return
            
            if act in ["write_file", "fs_edit"]:
                c_val = params.get("content", "") if act == "write_file" else params.get("new_string", "")
                if (params.get("filename", "").endswith(".py") or params.get("file_path", "").endswith(".py")) and c_val:
                    _, v_res = _validate_code({"code": c_val})
                    if not v_res.get("valid", True):
                        msgs.append({"role": "assistant", "content": json.dumps(dec)}); msgs.append({"role": "user", "content": f"Observation: CRITIC Error: {v_res.get('error')}"}); continue
            
            code, res = run_skill(act, params, caller)
            msgs.append({"role": "assistant", "content": json.dumps(dec)}); msgs.append({"role": "user", "content": f"Observation: {json.dumps({'skill': act, 'result': res})}"})
        
        with sqlite3.connect(DB_PATH) as conn:
            conn.execute("UPDATE jobs SET status = 'completed', result = ? WHERE task_id = ?", 
                         (json.dumps({"answer": "max iterations reached", "trace": trace}), task_id))
    except Exception as e:
        with sqlite3.connect(DB_PATH) as conn:
            conn.execute("UPDATE jobs SET status = 'failed', result = ? WHERE task_id = ?", 
                         (json.dumps({"error": str(e)}), task_id))

def run_skill(name, params, caller=None):
    skill = _catalog.get(name)
    if not skill:
        return 404, {"error": f"skill '{name}' not found"}
    if skill["upstream"] == "internal_exec":
        return _INTERNAL[name](params)
    path = skill["upstream"]
    for k, v in params.items():
        path = path.replace(f"{{{k}}}", str(v))
    try:
        with urllib.request.urlopen(f"{PROXY}{path}", timeout=5) as r:
            return r.status, json.loads(r.read())
    except urllib.error.HTTPError as e:
        return e.code, {"error": str(e.reason)}

# --- Servidor HTTP ---
class Handler(http.server.BaseHTTPRequestHandler):
    def log_message(self, *a): pass
    def send_json(self, code, data):
        body = json.dumps(data, ensure_ascii=False).encode()
        self.send_response(code); self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", len(body)); self.end_headers(); self.wfile.write(body)
    def do_POST(self):
        print(f"DEBUG: do_POST path={self.path}")
        length = int(self.headers.get("Content-Length", 0))
        body = json.loads(self.rfile.read(length)) if length else {}
        caller = self.headers.get("X-Delegation-Chain", "unknown")
        if caller == "unknown": return self.send_json(403, {"error": "missing X-Delegation-Chain"})

        # Endpoints que NÃO exigem 'prompt'
        if self.path in ["/run", "/tools/call"] or self.path.startswith("/acp/"):
            if self.path in ["/run", "/tools/call"]:
                skill_name = body.get("skill") if self.path == "/run" else body.get("name")
                params     = body.get("params", {}) if self.path == "/run" else body.get("arguments", {})
                if not skill_name: return self.send_json(400, {"error": "skill/name required"})
                code, result = run_skill(skill_name, params, caller)
                if self.path == "/tools/call":
                    return self.send_json(code, {"content": [{"type": "text", "text": json.dumps(result, ensure_ascii=False)}]})
                return self.send_json(code, {"caller": caller, "skill": skill_name, "result": result})

            # Handlers ACP
            if self.path == "/acp/discover":
                return self.send_json(200, {"agent_id": ACP_AGENT_ID, "capabilities": ACP_CAPABILITIES, "endpoints": ["/acp/negotiate", "/acp/execute"]})
            if self.path == "/acp/negotiate":
                requested = body.get("capabilities", [])
                available = [c for c in ACP_CAPABILITIES if c["name"] in requested]
                return self.send_json(200, {"agent_id": ACP_AGENT_ID, "supported": available, "unsupported": list(set(requested) - set(c["name"] for c in available))})
            if self.path == "/acp/execute":
                skill_name = body.get("skill")
                params = body.get("params", {})
                if not skill_name: return self.send_json(400, {"error": "skill required"})
                code, result = run_skill(skill_name, params, caller)
                return self.send_json(code, result)

        # Endpoint de CHAT (Exige 'prompt')
        prompt = body.get("prompt")
        if not prompt: return self.send_json(400, {"error": "prompt required"})
        task_id = str(uuid.uuid4())
        with sqlite3.connect(DB_PATH) as conn:
            conn.execute("INSERT INTO jobs (task_id, status, prompt) VALUES (?, 'queued', ?)", (task_id, prompt))
        threading.Thread(target=agent_worker, args=(task_id, prompt, caller), daemon=True).start()
        return self.send_json(202, {"task_id": task_id, "status": "queued"})
    
    def do_GET(self):
        parsed = urllib.parse.urlparse(self.path)
        if parsed.path == "/chat/status":
            query = urllib.parse.parse_qs(parsed.query); task_id = query.get("task_id", [None])[0]
            if not task_id: return self.send_json(400, {"error": "task_id required"})
            with sqlite3.connect(DB_PATH) as conn:
                row = conn.execute("SELECT status, result FROM jobs WHERE task_id = ?", (task_id,)).fetchone()
            if not row: return self.send_json(404, {"error": "not found"})
            res = {"task_id": task_id, "status": row[0]}
            if row[1]: res["result"] = json.loads(row[1])
            return self.send_json(200, res)
        if self.path == "/healthz": return self.send_json(200, {"status": "ok"})
        if self.path == "/skills": return self.send_json(200, list_skills())
        if self.path == "/tools":
            tools = [{"name": s["name"], "description": s["description"], "inputSchema": {"type": "object", "properties": {p: {"type": "string"} for p in s["params"]}, "required": s["params"]}, "systemInstruction": s.get("system_instruction", "")} for s in _catalog.values()]
            return self.send_json(200, {"schema": "mcp-tools/v1", "tools": tools})
        self.send_json(404, {"error": "not found"})

if __name__ == "__main__":
    init_db()
    # Instalação silenciosa da dependência RTI (Download direto via URL)
    import urllib.request, zipfile, io
    last30days_path = "/tmp/workspace/last30days-skill"
    if not os.path.exists(last30days_path):
        try:
            url = "https://github.com/mvanhorn/last30days-skill/archive/refs/heads/main.zip"
            with urllib.request.urlopen(url) as r:
                zip_data = io.BytesIO(r.read())
                with zipfile.ZipFile(zip_data) as z: z.extractall("/tmp/workspace")
            os.rename("/tmp/workspace/last30days-skill-main", last30days_path)
            subprocess.run(["python3", "-m", "pip", "install", "--user", "-r", os.path.join(last30days_path, "requirements.txt")], capture_output=True, check=False)
        except Exception as e: print(f"Erro na instalação RTI: {e}")

    server = http.server.HTTPServer(("0.0.0.0", 8080), Handler)
    server.socket.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
    print("LUDOC listening on :8080 (Async Motor + SQLite active)", flush=True)
    server.serve_forever()
